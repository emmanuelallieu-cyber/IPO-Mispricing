"""
=============================================================================
  Phase 5: SHAP Explainable AI (XAI) Analysis for IPO Underpricing Models
=============================================================================
  This script applies SHAP (SHapley Additive exPlanations) to the four ML
  models trained in Phase 3 & 4 to explain HOW and WHY each model makes its
  predictions on IPO underpricing.

  OUTPUT:
    - Multiple SHAP diagram PNGs (saved to working directory)
    - A consolidated Word Report: 'SHAP_XAI_IPO_Report.docx'

  REQUIREMENTS:
    pip install shap python-docx matplotlib seaborn scikit-learn xgboost lightgbm pandas numpy
=============================================================================
"""

import os
import re
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
import shap

from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from xgboost import XGBClassifier
from lightgbm import LGBMClassifier
from sklearn.metrics import accuracy_score, roc_auc_score
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore')
shap.initjs()

# ─────────────────────────────────────────────────────
#   HELPER: Apply dark navy header shading to a table row
# ─────────────────────────────────────────────────────
def shade_row(row, fill_hex="1F3864"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex="1F3864"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_image_centered(doc, path, width_inches=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    doc.add_paragraph()

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(80, 80, 80)

# ─────────────────────────────────────────────────────────
print("=" * 65)
print("   Phase 5: SHAP Explainable AI — IPO Underpricing Models")
print("=" * 65)

# ─────────────────────────────────────────────────────────
#   STEP 1 ─ Load & Prepare Data
# ─────────────────────────────────────────────────────────
print("\n[1/6] Loading and preparing data...")
try:
    X = pd.read_csv('X_Features_Engineered.csv')
    y = pd.read_csv('y_Target.csv').values.ravel()
except FileNotFoundError:
    print("\n  ERROR: 'X_Features_Engineered.csv' or 'y_Target.csv' not found.")
    print("  Please ensure Phase 3 outputs are in the same folder as this script.\n")
    raise

X = X.rename(columns=lambda c: re.sub('[^A-Za-z0-9_]+', '_', c))
feature_names = list(X.columns)

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)
X_train_df = pd.DataFrame(X_train, columns=feature_names)
X_test_df  = pd.DataFrame(X_test,  columns=feature_names)
print(f"  Data loaded: {X.shape[0]:,} samples · {X.shape[1]} features")
print(f"  Train: {len(X_train):,}  |  Test: {len(X_test):,}")

# ─────────────────────────────────────────────────────────
#   STEP 2 ─ Train All Four Models
# ─────────────────────────────────────────────────────────
print("\n[2/6] Training models...")
models = {
    "Logistic Regression": LogisticRegression(max_iter=2000, C=0.1, random_state=42),
    "Random Forest":       RandomForestClassifier(n_estimators=200, max_depth=10, random_state=42),
    "XGBoost":             XGBClassifier(n_estimators=100, learning_rate=0.05, eval_metric='logloss', random_state=42),
    "LightGBM":            LGBMClassifier(n_estimators=100, learning_rate=0.05, random_state=42, verbosity=-1),
}

metrics = {}
for name, model in models.items():
    model.fit(X_train_df, y_train)
    y_pred  = model.predict(X_test_df)
    y_probs = model.predict_proba(X_test_df)[:, 1]
    metrics[name] = {
        "accuracy": accuracy_score(y_test, y_pred),
        "auc":      roc_auc_score(y_test, y_probs),
    }
    print(f"  ✔  {name:<25}  Acc={metrics[name]['accuracy']:.2%}  AUC={metrics[name]['auc']:.4f}")

# Subsample for SHAP (speeds up computation; use full test set if small)
SHAP_SAMPLE = min(300, len(X_test_df))
X_shap = X_test_df.sample(n=SHAP_SAMPLE, random_state=42).reset_index(drop=True)

# ─────────────────────────────────────────────────────────
#   STEP 3 ─ Compute SHAP Values for Each Model
# ─────────────────────────────────────────────────────────
print(f"\n[3/6] Computing SHAP values (sample n={SHAP_SAMPLE})...")
explainers   = {}
shap_values  = {}

# --- Logistic Regression: LinearExplainer ---
print("  → Logistic Regression  (LinearExplainer)...")
exp_lr = shap.LinearExplainer(models["Logistic Regression"], X_train_df,
                               feature_perturbation="interventional")
sv_lr  = exp_lr(X_shap)
explainers["Logistic Regression"]  = exp_lr
shap_values["Logistic Regression"] = sv_lr

# --- Random Forest: TreeExplainer ---
print("  → Random Forest        (TreeExplainer)...")
exp_rf = shap.TreeExplainer(models["Random Forest"])
sv_rf  = exp_rf(X_shap, check_additivity=False)
# Keep only the positive-class SHAP values
if len(sv_rf.values.shape) == 3:
    sv_rf.values        = sv_rf.values[:, :, 1]
    sv_rf.base_values   = sv_rf.base_values[:, 1] if sv_rf.base_values.ndim > 1 else sv_rf.base_values
explainers["Random Forest"]  = exp_rf
shap_values["Random Forest"] = sv_rf

# --- XGBoost: TreeExplainer ---
print("  → XGBoost              (TreeExplainer)...")
exp_xgb = shap.TreeExplainer(models["XGBoost"])
sv_xgb  = exp_xgb(X_shap)
if len(sv_xgb.values.shape) == 3:
    sv_xgb.values      = sv_xgb.values[:, :, 1]
    sv_xgb.base_values = sv_xgb.base_values[:, 1] if sv_xgb.base_values.ndim > 1 else sv_xgb.base_values
explainers["XGBoost"]  = exp_xgb
shap_values["XGBoost"] = sv_xgb

# --- LightGBM: TreeExplainer ---
print("  → LightGBM             (TreeExplainer)...")
exp_lgb = shap.TreeExplainer(models["LightGBM"])
sv_lgb  = exp_lgb(X_shap)
if len(sv_lgb.values.shape) == 3:
    sv_lgb.values      = sv_lgb.values[:, :, 1]
    sv_lgb.base_values = sv_lgb.base_values[:, 1] if sv_lgb.base_values.ndim > 1 else sv_lgb.base_values
explainers["LightGBM"]  = exp_lgb
shap_values["LightGBM"] = sv_lgb

# ─────────────────────────────────────────────────────────
#   STEP 4 ─ Generate All SHAP Diagrams
# ─────────────────────────────────────────────────────────
print("\n[4/6] Generating SHAP diagrams...")
generated_plots = {}    # stores { model_name : { plot_type: filepath } }

COLORS = {
    "Logistic Regression": "#2980B9",
    "Random Forest":       "#27AE60",
    "XGBoost":             "#E67E22",
    "LightGBM":            "#8E44AD",
}

for name, sv in shap_values.items():
    safe = name.replace(" ", "_")
    plots = {}

    # ── (A) Beeswarm / Summary Plot ──────────────────────────────
    fname = f"SHAP_Beeswarm_{safe}.png"
    fig, ax = plt.subplots(figsize=(10, 7))
    shap.plots.beeswarm(sv, max_display=15, show=False)
    plt.title(f"SHAP Beeswarm — {name}\nTop 15 Features by Mean |SHAP|",
              fontsize=13, fontweight='bold', pad=14)
    plt.tight_layout()
    plt.savefig(fname, dpi=150, bbox_inches='tight')
    plt.close('all')
    plots['beeswarm'] = fname
    print(f"  ✔  {fname}")

    # ── (B) Bar Summary Plot ─────────────────────────────────────
    fname = f"SHAP_Bar_{safe}.png"
    fig, ax = plt.subplots(figsize=(10, 7))
    shap.plots.bar(sv, max_display=15, show=False)
    plt.title(f"SHAP Feature Importance (Bar) — {name}",
              fontsize=13, fontweight='bold', pad=14)
    plt.tight_layout()
    plt.savefig(fname, dpi=150, bbox_inches='tight')
    plt.close('all')
    plots['bar'] = fname
    print(f"  ✔  {fname}")

    # ── (C) Waterfall for Best Predicted IPO ─────────────────────
    fname = f"SHAP_Waterfall_{safe}.png"
    probs    = models[name].predict_proba(X_shap)[:, 1]
    best_idx = int(np.argmax(probs))           # highest predicted probability

    # Extract SHAP values for this single observation and guarantee
    # base_values is a plain Python float (waterfall requires 0-d scalar).
    single_shap_vals = sv.values[best_idx]     # shape (n_features,)
    raw_base = sv.base_values[best_idx]
    base_scalar = float(np.asarray(raw_base).ravel()[0])  # handles 0-d, 1-d, 2-d arrays
    single_data  = sv.data[best_idx] if sv.data is not None else X_shap.iloc[best_idx].values

    single_exp = shap.Explanation(
        values      = single_shap_vals,
        base_values = base_scalar,
        data        = single_data,
        feature_names = feature_names,
    )

    fig, ax = plt.subplots(figsize=(10, 7))
    shap.plots.waterfall(single_exp, max_display=12, show=False)
    plt.title(f"SHAP Waterfall — {name}\nHighest-Confidence Underpriced IPO Prediction",
              fontsize=12, fontweight='bold', pad=14)
    plt.tight_layout()
    plt.savefig(fname, dpi=150, bbox_inches='tight')
    plt.close('all')
    plots['waterfall'] = fname
    print(f"  ✔  {fname}")

    # ── (D) Mean |SHAP| Top-10 Custom Bar (branded) ──────────────
    fname = f"SHAP_MeanAbs_{safe}.png"
    mean_abs = pd.Series(
        np.abs(sv.values).mean(axis=0), index=feature_names
    ).sort_values(ascending=True).tail(10)

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.barh(mean_abs.index, mean_abs.values,
                   color=COLORS[name], edgecolor='white', height=0.65)
    ax.bar_label(bars, fmt="%.4f", padding=4, fontsize=9)
    ax.set_xlabel("Mean |SHAP Value|", fontsize=11)
    ax.set_title(f"Top-10 Mean Absolute SHAP — {name}", fontsize=13, fontweight='bold')
    ax.spines[['top','right']].set_visible(False)
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('white')
    plt.tight_layout()
    plt.savefig(fname, dpi=150, bbox_inches='tight')
    plt.close('all')
    plots['mean_abs'] = fname
    print(f"  ✔  {fname}")

    generated_plots[name] = plots

# ── (E) Cross-Model Feature Importance Heatmap ───────────────────
print("  → Cross-model comparison heatmap...")
fname_heat = "SHAP_CrossModel_Heatmap.png"
all_mean_abs = {}
for name, sv in shap_values.items():
    all_mean_abs[name] = pd.Series(np.abs(sv.values).mean(axis=0), index=feature_names)
heat_df = pd.DataFrame(all_mean_abs)
# Rank by total importance across models
heat_df['total'] = heat_df.sum(axis=1)
top20 = heat_df.sort_values('total', ascending=False).head(20).drop(columns='total')
# Normalize each model column 0→1
norm_df = top20.apply(lambda c: (c - c.min()) / (c.max() - c.min() + 1e-9))

fig, ax = plt.subplots(figsize=(11, 9))
sns.heatmap(norm_df, annot=True, fmt=".2f", cmap="YlOrRd",
            linewidths=0.5, ax=ax, cbar_kws={'label': 'Normalised Mean |SHAP|'})
ax.set_title("Cross-Model SHAP Feature Importance (Top 20 Features)\nNormalised per Model",
             fontsize=13, fontweight='bold', pad=14)
ax.set_xlabel("Model", fontsize=11)
ax.set_ylabel("Feature", fontsize=11)
plt.tight_layout()
plt.savefig(fname_heat, dpi=150, bbox_inches='tight')
plt.close('all')
print(f"  ✔  {fname_heat}")

# ── (F) SHAP Dependence Plot for top shared feature ──────────────
shared_top = top20.index[0]
print(f"  → SHAP dependence plots (top shared feature: '{shared_top}')...")
dep_files = {}
for name, sv in shap_values.items():
    safe = name.replace(" ", "_")
    fname = f"SHAP_Dependence_{safe}.png"
    feat_idx = feature_names.index(shared_top)
    fig, ax = plt.subplots(figsize=(9, 5))
    sc = ax.scatter(
        X_shap.iloc[:, feat_idx],
        sv.values[:, feat_idx],
        c=sv.values[:, feat_idx],
        cmap='coolwarm', alpha=0.6, edgecolors='none', s=25
    )
    plt.colorbar(sc, ax=ax, label='SHAP Value')
    ax.axhline(0, color='grey', lw=0.8, ls='--')
    ax.set_xlabel(shared_top, fontsize=11)
    ax.set_ylabel(f"SHAP Value for '{shared_top}'", fontsize=11)
    ax.set_title(f"SHAP Dependence — {name}\nFeature: {shared_top}", fontsize=12, fontweight='bold')
    ax.spines[['top','right']].set_visible(False)
    plt.tight_layout()
    plt.savefig(fname, dpi=150, bbox_inches='tight')
    plt.close('all')
    dep_files[name] = fname
    generated_plots[name]['dependence'] = fname
    print(f"  ✔  {fname}")

# ─────────────────────────────────────────────────────────
#   STEP 5 ─ Build the Word Document Report
# ─────────────────────────────────────────────────────────
print("\n[5/6] Building Word document report...")

doc = Document()

# ── Page margins (1 inch all round) ──
section = doc.sections[0]
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)

# ═══════════════════════════════════════════════════════
#   COVER PAGE
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("SHAP Explainable AI Report")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = RGBColor(31, 56, 100)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Machine Learning Models for IPO Underpricing Prediction")
sr.font.size = Pt(16)
sr.font.color.rgb = RGBColor(70, 70, 70)

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Phase 5 — Explainability & Interpretability Analysis",
    "Method: SHapley Additive exPlanations (SHAP)",
    "Models: Logistic Regression · Random Forest · XGBoost · LightGBM",
]:
    info_p.add_run(line + "\n").font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#   SECTION 1 — EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════
add_heading(doc, "1. Executive Summary", level=1)
add_para(doc,
    "This report presents an in-depth Explainable AI (XAI) analysis of four machine learning models "
    "trained to predict IPO underpricing. While Phase 3 & 4 of this project established that tree-based "
    "ensemble models (Random Forest, XGBoost, LightGBM) outperform linear baselines in discriminative "
    "power, a black-box model alone is insufficient for investment-grade decision support. This report "
    "applies SHAP — the gold standard framework in model interpretability — to reveal the internal "
    "logic of each model, quantify each feature's contribution, and ensure the predictions are "
    "trustworthy, explainable, and regulatorily defensible."
)

# ─── Performance summary table ────────────────────────
add_heading(doc, "1.1 Model Performance Summary", level=2)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0]
for i, txt in enumerate(["Model", "Accuracy", "AUC Score"]):
    cell = hdr.cells[i]
    cell.text = txt
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_row(hdr, "1F3864")

row_fills = ["EAF1FB", "FFFFFF"]
for i, (name, m) in enumerate(metrics.items()):
    row = tbl.add_row()
    vals = [name, f"{m['accuracy']:.2%}", f"{m['auc']:.4f}"]
    for j, v in enumerate(vals):
        row.cells[j].text = v
        row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(row.cells[j], row_fills[i % 2])
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#   SECTION 2 — WHAT IS SHAP
# ═══════════════════════════════════════════════════════
add_heading(doc, "2. Methodology: SHapley Additive exPlanations (SHAP)", level=1)
add_para(doc,
    "SHAP is a game-theory-based framework that assigns each feature a 'Shapley value' — the average "
    "marginal contribution of that feature across all possible combinations of features. Unlike simple "
    "feature importance scores, SHAP values are:"
)
bullets = [
    "Consistent — a feature that increases a model's output is always assigned a higher SHAP value.",
    "Local & Global — explain both individual predictions (local) and overall model behaviour (global).",
    "Signed — positive SHAP values push the prediction toward underpricing; negative values push away.",
    "Additive — the sum of all SHAP values for a prediction equals the deviation from the model's base rate.",
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(b).font.size = Pt(11)

doc.add_paragraph()
add_para(doc,
    "Each model class uses the most efficient SHAP explainer algorithm for its structure:"
)
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(["Model", "SHAP Explainer", "Complexity"]):
    cell = tbl2.rows[0].cells[i]
    cell.text = txt
    for run in cell.paragraphs[0].runs:
        run.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_row(tbl2.rows[0], "1F3864")
explainer_info = [
    ("Logistic Regression", "LinearExplainer",  "O(M) — Linear algebra, exact"),
    ("Random Forest",       "TreeExplainer",     "O(T·L²) — Polynomial"),
    ("XGBoost",             "TreeExplainer",     "O(T·D) — Fast tree traversal"),
    ("LightGBM",            "TreeExplainer",     "O(T·D) — Fast tree traversal"),
]
for i, row_data in enumerate(explainer_info):
    row = tbl2.add_row()
    for j, v in enumerate(row_data):
        row.cells[j].text = v
        row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(row.cells[j], "EAF1FB" if i % 2 == 0 else "FFFFFF")
doc.add_paragraph()

add_para(doc,
    f"For computational efficiency, SHAP values were computed on a random subsample of "
    f"{SHAP_SAMPLE} test-set observations. This is standard practice; SHAP's theoretical guarantees "
    "hold on any representative subsample."
)

# ═══════════════════════════════════════════════════════
#   SECTION 3 — CHART TYPE GUIDE
# ═══════════════════════════════════════════════════════
add_heading(doc, "3. How to Read SHAP Diagrams", level=1)

chart_types = [
    ("Beeswarm Plot",
     "Each dot is one test observation. The x-axis shows the SHAP value (impact on prediction). "
     "Colour indicates the raw feature value (red = high, blue = low). Features are ranked top-to-bottom "
     "by mean absolute SHAP value. A cluster to the right means high feature values push toward predicting "
     "underpricing; a cluster to the left means they push away."),
    ("Bar Plot",
     "Shows the global average of absolute SHAP values per feature. This is the single most reliable "
     "measure of overall feature importance — larger bars mean the model leans more heavily on that feature."),
    ("Waterfall Plot",
     "Explains one specific IPO prediction. It starts at the model's base rate (E[f(x)]) and shows how "
     "each feature pushes the final prediction up (red) or down (blue) from that baseline. The result is "
     "the model's predicted probability for that individual IPO."),
    ("Mean |SHAP| Bar Chart",
     "A custom clean version of the Bar Plot showing the top-10 features with numerical labels. "
     "Useful for quick executive communication."),
    ("Cross-Model Heatmap",
     "Compares feature importance across all four models simultaneously. Each cell is the normalised "
     "mean absolute SHAP value (0 = least important, 1 = most important for that model). Features "
     "consistently dark across all columns are the most robust predictors of IPO underpricing."),
    ("Dependence Plot",
     "Scatter plot showing how the SHAP value of the top shared feature changes with its actual value. "
     "Colour encodes the SHAP value itself. Reveals non-linear relationships, thresholds, and "
     "interaction effects that simpler importance metrics miss."),
]
for title, desc in chart_types:
    p = doc.add_paragraph()
    p.add_run(f"  {title}: ").bold = True
    p.add_run(desc).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#   SECTION 4 — PER-MODEL ANALYSIS
# ═══════════════════════════════════════════════════════
add_heading(doc, "4. Per-Model SHAP Analysis", level=1)

model_descriptions = {
    "Logistic Regression": (
        "Logistic Regression is a linear probabilistic classifier. Its SHAP values are directly "
        "proportional to the model's regression coefficients multiplied by the input feature values "
        "(after centring). Because the model is constrained to linear relationships, SHAP values "
        "reveal a clean, additive story: each feature independently contributes a fixed amount per "
        "unit of its value. This makes Logistic Regression the most interpretable baseline and a "
        "regulatory-compliance favourite."
    ),
    "Random Forest": (
        "Random Forest aggregates predictions from 200 decision trees, each trained on a bootstrap "
        "sample with random feature subsets. SHAP uses the TreeExplainer algorithm, which exploits "
        "the tree structure to compute exact Shapley values in polynomial time. The result captures "
        "feature interactions and non-linearities that Logistic Regression cannot model — at the cost "
        "of reduced transparency per-tree."
    ),
    "XGBoost": (
        "XGBoost is a gradient-boosted tree ensemble. Unlike Random Forest's parallel trees, XGBoost "
        "builds trees sequentially, each correcting the residuals of the previous. TreeExplainer "
        "computes exact SHAP values by traversing each tree's decision paths. XGBoost's SHAP profile "
        "tends to be sparser (fewer features drive most of the prediction) compared to Random Forest, "
        "reflecting its regularisation via shrinkage and tree depth constraints."
    ),
    "LightGBM": (
        "LightGBM is Microsoft's gradient-boosted framework optimised for speed via histogram-based "
        "binning and leaf-wise tree growth. Its SHAP values are computed identically to XGBoost via "
        "TreeExplainer. LightGBM's leaf-wise strategy allows deeper, more asymmetric trees, which can "
        "capture sharp threshold effects invisible to depth-limited models like XGBoost with "
        "max_depth constraints."
    ),
}

for name, plots in generated_plots.items():
    safe = name.replace(" ", "_")
    color = COLORS[name]
    r, g, b = int(color[1:3],16), int(color[3:5],16), int(color[5:7],16)

    add_heading(doc, f"4.{list(generated_plots.keys()).index(name)+1}  {name}", level=2)
    
    # coloured divider line via paragraph border
    p_div = doc.add_paragraph()
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color[1:])
    pBdr.append(bottom)
    pPr.append(pBdr)
    p_div.paragraph_format.space_after = Pt(4)

    add_para(doc, model_descriptions[name])
    add_para(doc, f"  Performance — Accuracy: {metrics[name]['accuracy']:.2%}   AUC: {metrics[name]['auc']:.4f}",
             bold=True)

    # Beeswarm
    add_heading(doc, "Beeswarm Plot (Global Feature Impact)", level=3)
    add_para(doc,
        "The beeswarm diagram below shows the distribution of SHAP values for the top 15 features "
        "across all test-set IPOs. The x-axis encodes direction and magnitude of impact. Features at "
        "the top have the highest average influence on predictions.")
    if os.path.exists(plots['beeswarm']):
        add_image_centered(doc, plots['beeswarm'], width_inches=5.8)
        add_caption(doc, f"Figure: SHAP Beeswarm Plot — {name}. "
                        "Red dots = high feature value; Blue dots = low feature value.")

    # Bar chart
    add_heading(doc, "Mean |SHAP| Feature Importance", level=3)
    add_para(doc,
        "The bar chart ranks features by their average absolute SHAP contribution, providing an "
        "unambiguous global importance ranking. Unlike traditional feature importances (e.g. Gini "
        "impurity in Random Forest), SHAP-based importance is consistent across model architectures.")
    if os.path.exists(plots['mean_abs']):
        add_image_centered(doc, plots['mean_abs'], width_inches=5.8)
        add_caption(doc, f"Figure: Top-10 Mean |SHAP| Values — {name}.")

    # Waterfall
    add_heading(doc, "Waterfall Plot (Individual Prediction)", level=3)
    add_para(doc,
        "The waterfall plot explains the model's highest-confidence underpricing prediction in detail. "
        "Starting from the base rate E[f(x)], each bar shows how a feature moves the prediction. "
        "Red bars push toward predicting underpricing (positive class); blue bars push away. "
        "The final value f(x) is the model's predicted probability for this IPO.")
    if os.path.exists(plots['waterfall']):
        add_image_centered(doc, plots['waterfall'], width_inches=5.8)
        add_caption(doc, f"Figure: SHAP Waterfall Plot — {name}. "
                        "Highest-confidence underpriced IPO from the test set.")

    # Dependence
    add_heading(doc, f"Dependence Plot — '{shared_top}'", level=3)
    add_para(doc,
        f"The dependence plot reveals the relationship between the feature '{shared_top}' "
        "and its SHAP contribution. A non-horizontal pattern signals a non-linear effect: the feature "
        "has a disproportionate impact at certain value thresholds, which is critical for identifying "
        "high-risk IPO characteristics.")
    if 'dependence' in plots and os.path.exists(plots['dependence']):
        add_image_centered(doc, plots['dependence'], width_inches=5.5)
        add_caption(doc, f"Figure: SHAP Dependence Plot — {name}, Feature: '{shared_top}'.")

    doc.add_page_break()

# ═══════════════════════════════════════════════════════
#   SECTION 5 — CROSS-MODEL COMPARISON
# ═══════════════════════════════════════════════════════
add_heading(doc, "5. Cross-Model Feature Importance Comparison", level=1)
add_para(doc,
    "The heatmap below compares the normalised mean absolute SHAP values across all four models "
    "for the top 20 most important features globally. This comparative view is essential for "
    "identifying features that are robust predictors regardless of model architecture — such "
    "features represent the most reliable IPO underpricing signals in the dataset."
)
add_para(doc,
    "Features that are consistently dark (high) across all four columns should be prioritised in "
    "investment analysis workflows. Features that are important in only one or two models may "
    "reflect model-specific artefacts or interaction effects unique to that algorithm's inductive bias."
)
if os.path.exists(fname_heat):
    add_image_centered(doc, fname_heat, width_inches=6.2)
    add_caption(doc, "Figure: Cross-Model SHAP Heatmap. "
                     "Darker cells = higher relative feature importance for that model. "
                     "Values are normalised 0–1 per column.")

doc.add_paragraph()
add_para(doc, "Key observations from the cross-model analysis:", bold=True)
obs_bullets = [
    "Features consistently important across tree models (RF, XGBoost, LightGBM) but not in "
    "Logistic Regression indicate non-linear predictive relationships.",
    "Features prominent only in Logistic Regression capture linear signals that tree models may "
    "under-utilise due to splits favouring non-linear boundaries.",
    "Perfect alignment across all four models provides the strongest evidence that a feature "
    "is a fundamental driver of IPO underpricing, not a modelling artefact.",
]
for b in obs_bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(b).font.size = Pt(11)

# ═══════════════════════════════════════════════════════
#   SECTION 6 — IMPLICATIONS FOR IPO ANALYSIS
# ═══════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "6. Implications for IPO Underpricing Analysis", level=1)
add_para(doc,
    "The SHAP analysis transforms the models from opaque prediction engines into interpretable "
    "decision-support tools. Below are the principal actionable insights:"
)

implications = [
    ("Feature Engineering Validation",
     "SHAP confirms which engineered features carry genuine predictive signal vs. which are noise. "
     "Features with near-zero SHAP values across all models should be candidates for removal in "
     "future pipeline iterations to reduce overfitting risk."),
    ("Non-Linearity Detection",
     "SHAP dependence plots expose threshold effects — specific ranges of a feature value where "
     "underpricing probability surges. These thresholds can be translated directly into investment "
     "screening criteria (e.g., 'flag any IPO where Feature X exceeds threshold T')."),
    ("Model Selection Guidance",
     "If the top SHAP features are consistent across all four models, any model can be deployed "
     "with confidence. If feature rankings diverge substantially, the ensemble approach (or further "
     "model validation) is recommended before live deployment."),
    ("Regulatory & Audit Readiness",
     "For institutional investors, SHAP-powered explanations can satisfy MiFID II / internal "
     "governance requirements by providing an auditable, feature-level justification for each "
     "model-driven investment decision."),
    ("Ongoing Monitoring",
     "SHAP value distributions can be tracked over time. A significant shift in which features "
     "drive predictions signals concept drift — an early warning that the model needs retraining "
     "on more recent IPO data."),
]
for title, desc in implications:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(f"  {title}: ").bold = True
    p.add_run(desc).font.size = Pt(11)

# ═══════════════════════════════════════════════════════
#   SECTION 7 — CONCLUSION
# ═══════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "7. Conclusion", level=1)
add_para(doc,
    "This report has demonstrated that SHAP-based Explainable AI is not merely a post-hoc "
    "visualisation exercise but a core component of responsible machine learning deployment "
    "in financial applications. By computing exact Shapley values for all four IPO underpricing "
    "models, we have:"
)
conclusions = [
    "Identified the top features driving underpricing predictions with mathematical rigour.",
    "Explained individual predictions at the level of single IPOs — enabling case-by-case scrutiny.",
    "Compared model behaviour across four architectures to surface the most robust signals.",
    "Detected potential non-linearities and threshold effects invisible to standard importance metrics.",
    "Produced an audit-ready report suitable for investment committee and compliance review.",
]
for c in conclusions:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c).font.size = Pt(11)

doc.add_paragraph()
add_para(doc,
    "The next recommended step is to integrate SHAP-based monitoring into the model's production "
    "pipeline so that prediction explanations are generated automatically for every new IPO scored, "
    "ensuring continuous transparency throughout the model's operational lifetime."
)

# ── Footer with page numbers ──────────────────────────
for section in doc.sections:
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("SHAP XAI Analysis — IPO Underpricing Prediction | Confidential").font.size = Pt(9)

# ─────────────────────────────────────────────────────────
#   STEP 6 ─ Save & Report
# ─────────────────────────────────────────────────────────
docx_path = "SHAP_XAI_IPO_Report.docx"
doc.save(docx_path)
print(f"\n[6/6] Word report saved → '{docx_path}'")

print("\n" + "=" * 65)
print("  SHAP XAI Analysis Complete!")
print("=" * 65)
print("\n  Generated Files:")
all_files = []
for name, plots in generated_plots.items():
    for ptype, fpath in plots.items():
        if os.path.exists(fpath):
            all_files.append(fpath)
            print(f"    PNG  {fpath}")
if os.path.exists(fname_heat):
    print(f"    PNG  {fname_heat}")
print(f"    DOCX {docx_path}")
print("\n  Open 'SHAP_XAI_IPO_Report.docx' for the complete analysis.\n")
